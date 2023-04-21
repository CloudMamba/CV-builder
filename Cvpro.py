from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)
document = Document()

#profile picture
document.add_picture('evening.jpg', 
width = Inches(2.0))

#name,phone number and email details
name = input("what is your name?")
speak('Hello'+ name + 'how are you today')

speak('What is your phone number?')
phone_number = input("what is your phone number?")
email = input("what is your email?")

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)
#about me
document.add_heading('About me')
document.add_paragraph(
    input("Tell me about yourself:\n"))

#work experience

document.add_heading("Work Experience")
p = document.add_paragraph()

company = input('Enter Company:\n')
from_date = input ('From Date:\n')
to_date = input("To date:\n")

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-'+ to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

#more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No\n'
    )
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company:\n')
        from_date = input ('From Date:\n')
        to_date = input("To date:\n")

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-'+ to_date + '\n').italic = True

        experience_details = input(
          'Describe your experience at ' + company + ':\n')
        p.add_run(experience_details)
    else:
        break


#skills
document.add_heading('Skills')
Skill = input('Enter skill\n')
p = document.add_paragraph(Skill)
p.style = 'List Bullet'

#more skills
while True:
    has_more_skills = input(
        "Do you have more skills?, Yes or No\n")
    if has_more_skills.lower() == 'yes':
        Skill = input('Input your skill\n')
        p = document.add_paragraph(Skill) 
        p.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraph[0]
p.text = 'CV  generated using python language with the docx library'

document.save('cv.docx') 